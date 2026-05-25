#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
D&D 5e Character Sheet Generator with Auto-Calculation
Генератор листа персонажа D&D 5e с авторасчётом и экспортом в DOCX

Features:
- All fields from the standard D&D 5e character sheet (3 pages)
- Auto-calculation of ability modifiers, saving throws, skills,
  proficiency bonus, initiative, passive perception, spell stats
- Export to readable, editable DOCX format
- Long text fields (backstory, etc.) expand as needed

Usage:
  python3 dnd_character_sheet_generator.py

Author: Z.ai
"""

import math
import json
import os
from dataclasses import dataclass, field, asdict
from typing import Optional, List, Dict, Tuple

# ─── Auto-calculation Engine ───────────────────────────────────────────────────

ABILITY_NAMES = ['СИЛ', 'ЛОВ', 'ТЕЛ', 'ИНТ', 'МДР', 'ХАР']
ABILITY_FULL_NAMES = {
    'СИЛ': 'Сила',
    'ЛОВ': 'Ловкость',
    'ТЕЛ': 'Телосложение',
    'ИНТ': 'Интеллект',
    'МДР': 'Мудрость',
    'ХАР': 'Харизма',
}
ABILITY_ABBR_TO_ENG = {
    'СИЛ': 'STR', 'ЛОВ': 'DEX', 'ТЕЛ': 'CON',
    'ИНТ': 'INT', 'МДР': 'WIS', 'ХАР': 'CHA',
}

SKILL_MAP = {
    'Акробатика': 'ЛОВ',
    'Анализ': 'ИНТ',
    'Атлетика': 'СИЛ',
    'Внимательность': 'МДР',
    'Выживание': 'МДР',
    'Выступление': 'ХАР',
    'Запугивание': 'ХАР',
    'История': 'ИНТ',
    'Ловкость рук': 'ЛОВ',
    'Магия': 'ИНТ',
    'Медицина': 'МДР',
    'Обман': 'ХАР',
    'Природа': 'ИНТ',
    'Проницательность': 'МДР',
    'Религия': 'ИНТ',
    'Скрытность': 'ЛОВ',
    'Убеждение': 'ХАР',
    'Уход за животными': 'МДР',
}

ALL_SKILLS = list(SKILL_MAP.keys())


def calc_modifier(score: int) -> int:
    """Calculate ability modifier from score."""
    return math.floor((score - 10) / 2)


def calc_proficiency_bonus(level: int) -> int:
    """Calculate proficiency bonus based on character level."""
    if level < 1:
        return 2
    if level <= 4:
        return 2
    elif level <= 8:
        return 3
    elif level <= 12:
        return 4
    elif level <= 16:
        return 5
    else:
        return 6


def format_modifier(mod: int) -> str:
    """Format modifier with sign."""
    return f"+{mod}" if mod >= 0 else str(mod)


# ─── Character Data Model ─────────────────────────────────────────────────────

@dataclass
class Attack:
    """Attack/spell entry."""
    name: str = ""
    attack_bonus: str = ""  # Can be auto or manual
    damage_and_type: str = ""


@dataclass
class SpellSlotInfo:
    """Spell slot information for a level."""
    level: int = 0
    total_slots: int = 0
    expended_slots: int = 0


@dataclass
class SpellEntry:
    """A single spell entry."""
    name: str = ""
    prepared: bool = False


@dataclass
class SpellLevelSpells:
    """Spells grouped by level."""
    level: int = 0  # 0 = cantrips
    spells: List[SpellEntry] = field(default_factory=list)


@dataclass
class Character:
    """Complete D&D 5e Character data model."""

    # ── Basic Info ──
    name: str = ""
    class_name: str = ""
    level: int = 1
    background: str = ""
    player_name: str = ""
    race: str = ""
    alignment: str = ""
    experience_points: int = 0
    inspiration: bool = False

    # ── Ability Scores (base, before racial/other bonuses) ──
    ability_scores: Dict[str, int] = field(default_factory=lambda: {
        'СИЛ': 10, 'ЛОВ': 10, 'ТЕЛ': 10,
        'ИНТ': 10, 'МДР': 10, 'ХАР': 10
    })

    # ── Racial / Other bonuses to abilities ──
    ability_bonuses: Dict[str, int] = field(default_factory=lambda: {
        'СИЛ': 0, 'ЛОВ': 0, 'ТЕЛ': 0,
        'ИНТ': 0, 'МДР': 0, 'ХАР': 0
    })

    # ── Saving Throw Proficiencies ──
    saving_throw_proficiencies: Dict[str, bool] = field(default_factory=lambda: {
        'СИЛ': False, 'ЛОВ': False, 'ТЕЛ': False,
        'ИНТ': False, 'МДР': False, 'ХАР': False
    })

    # ── Skill Proficiencies ──
    skill_proficiencies: Dict[str, bool] = field(default_factory=dict)

    # ── Skill Expertise (double proficiency) ──
    skill_expertise: Dict[str, bool] = field(default_factory=dict)

    # ── Combat Stats (some auto-calculated, some manual overrides) ──
    armor_class: Optional[int] = None  # None = auto from DEX if unarmored, or manual
    initiative_override: Optional[int] = None  # None = auto (DEX mod)
    speed: int = 30
    hp_max: Optional[int] = None  # None = auto-calculate
    hp_current: int = 0
    hp_temp: int = 0
    hit_dice: str = ""  # e.g., "1d10"

    # ── Death Saves ──
    death_save_successes: int = 0
    death_save_failures: int = 0

    # ── Attacks & Spellcasting ──
    attacks: List[Attack] = field(default_factory=list)

    # ── Currency (in pieces) ──
    cp: int = 0  # ММ - медные
    sp: int = 0  # СМ - серебряные
    ep: int = 0  # ЭМ - электрумовые
    gp: int = 0  # ЗМ - золотые
    pp: int = 0  # ПМ - платиновые

    # ── Personality ──
    personality_traits: str = ""
    ideals: str = ""
    bonds: str = ""
    flaws: str = ""

    # ── Other Proficiencies & Languages ──
    other_proficiencies_languages: str = ""

    # ── Features & Traits ──
    features_traits: str = ""

    # ── Equipment ──
    equipment: str = ""

    # ── Page 2: Character Details ──
    age: str = ""
    height: str = ""
    weight: str = ""
    eyes: str = ""
    skin: str = ""
    hair: str = ""
    appearance: str = ""
    allies_organizations: str = ""
    additional_features_traits: str = ""
    backstory: str = ""
    treasure: str = ""

    # ── Page 3: Spellcasting ──
    spellcasting_class: str = ""
    spellcasting_ability: str = ""  # СИЛ/ЛОВ/ТЕЛ/ИНТ/МДР/ХАР
    spell_save_dc_override: Optional[int] = None
    spell_attack_bonus_override: Optional[int] = None

    # Spell slots per level (1-9)
    spell_slots: Dict[int, SpellSlotInfo] = field(default_factory=dict)

    # Cantrips
    cantrips: List[str] = field(default_factory=list)

    # Spells by level
    spells_by_level: Dict[int, List[SpellEntry]] = field(default_factory=dict)

    # ── Methods ──

    def get_total_ability_score(self, ability: str) -> int:
        """Get total ability score including bonuses."""
        return self.ability_scores.get(ability, 10) + self.ability_bonuses.get(ability, 0)

    def get_modifier(self, ability: str) -> int:
        """Get ability modifier."""
        return calc_modifier(self.get_total_ability_score(ability))

    def get_proficiency_bonus(self) -> int:
        """Get proficiency bonus based on level."""
        return calc_proficiency_bonus(self.level)

    def get_saving_throw(self, ability: str) -> int:
        """Get saving throw bonus."""
        mod = self.get_modifier(ability)
        if self.saving_throw_proficiencies.get(ability, False):
            mod += self.get_proficiency_bonus()
        return mod

    def get_skill_bonus(self, skill: str) -> int:
        """Get skill bonus."""
        ability = SKILL_MAP.get(skill, 'СИЛ')
        mod = self.get_modifier(ability)
        if self.skill_expertise.get(skill, False):
            mod += self.get_proficiency_bonus() * 2
        elif self.skill_proficiencies.get(skill, False):
            mod += self.get_proficiency_bonus()
        return mod

    def get_initiative(self) -> int:
        """Get initiative bonus."""
        if self.initiative_override is not None:
            return self.initiative_override
        return self.get_modifier('ЛОВ')

    def get_passive_perception(self) -> int:
        """Get passive perception score."""
        mod = self.get_modifier('МДР')
        if self.skill_proficiencies.get('Внимательность', False):
            mod += self.get_proficiency_bonus()
        return 10 + mod

    def get_spell_save_dc(self) -> int:
        """Get spell save DC."""
        if self.spell_save_dc_override is not None:
            return self.spell_save_dc_override
        if not self.spellcasting_ability:
            return 0
        return 8 + self.get_proficiency_bonus() + self.get_modifier(self.spellcasting_ability)

    def get_spell_attack_bonus(self) -> int:
        """Get spell attack bonus."""
        if self.spell_attack_bonus_override is not None:
            return self.spell_attack_bonus_override
        if not self.spellcasting_ability:
            return 0
        return self.get_proficiency_bonus() + self.get_modifier(self.spellcasting_ability)

    def get_spellcasting_ability_modifier(self) -> int:
        """Get spellcasting ability modifier."""
        if not self.spellcasting_ability:
            return 0
        return self.get_modifier(self.spellcasting_ability)

    def get_hp_max(self) -> int:
        """Get max HP (auto or manual)."""
        if self.hp_max is not None:
            return self.hp_max
        return 0  # Cannot auto-calculate without class hit die + CON

    def init_skill_defaults(self):
        """Initialize skill proficiency dicts with defaults."""
        for skill in ALL_SKILLS:
            if skill not in self.skill_proficiencies:
                self.skill_proficiencies[skill] = False
            if skill not in self.skill_expertise:
                self.skill_expertise[skill] = False

    def to_dict(self) -> dict:
        """Convert to dictionary for JSON export."""
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict) -> 'Character':
        """Create Character from dictionary."""
        # Handle nested objects
        attacks = [Attack(**a) if isinstance(a, dict) else a for a in data.get('attacks', [])]
        spell_slots = {}
        for k, v in data.get('spell_slots', {}).items():
            spell_slots[int(k)] = SpellSlotInfo(**v) if isinstance(v, dict) else v
        spells_by_level = {}
        for k, v in data.get('spells_by_level', {}).items():
            level = int(k)
            spells_by_level[level] = [SpellEntry(**s) if isinstance(s, dict) else s for s in v]

        char = cls()
        for key, value in data.items():
            if key == 'attacks':
                char.attacks = attacks
            elif key == 'spell_slots':
                char.spell_slots = spell_slots
            elif key == 'spells_by_level':
                char.spells_by_level = spells_by_level
            else:
                setattr(char, key, value)
        char.init_skill_defaults()
        return char


# ─── DOCX Generator ───────────────────────────────────────────────────────────

def generate_docx(character: Character, output_path: str):
    """Generate a formatted DOCX character sheet."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Color constants
    COLOR_HEADER_BG = "2C3E50"
    COLOR_SUBHEADER_BG = "34495E"
    COLOR_LIGHT_BG = "ECF0F1"
    COLOR_VERY_LIGHT_BG = "F8F9FA"
    COLOR_ACCENT = "2980B9"
    COLOR_BORDER = "BDC3C7"
    COLOR_CALCULATED = "8E44AD"  # Purple for auto-calculated values
    COLOR_TEXT = "2C3E50"

    def set_cell_shading(cell, color: str):
        """Set cell background color."""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_border(cell, **kwargs):
        """Set cell borders."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>')
        for edge, val in kwargs.items():
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
                f'w:sz="{val.get("sz", "4")}" w:space="0" '
                f'w:color="{val.get("color", "000000")}"/>'
            )
            tcBorders.append(element)
        tcPr.append(tcBorders)

    def add_cell_text(cell, text: str, bold: bool = False, size: int = 9,
                      color: str = COLOR_TEXT, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      italic: bool = False):
        """Add formatted text to a table cell."""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.italic = italic
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def add_cell_multiline(cell, lines: list, size: int = 9, color: str = COLOR_TEXT,
                           bold_first: bool = False):
        """Add multiple lines to a cell."""
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        for i, line in enumerate(lines):
            if i > 0:
                run = p.add_run("\n")
                run.font.size = Pt(size)
            run = p.add_run(line)
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(color)
            if bold_first and i == 0:
                run.font.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def add_section_header(doc, text: str):
        """Add a section header paragraph."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLOR_HEADER_BG)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="{COLOR_ACCENT}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def make_header_row(table, row_idx, texts, bg_color=COLOR_HEADER_BG):
        """Format a row as a header row."""
        for i, text in enumerate(texts):
            cell = table.cell(row_idx, i)
            set_cell_shading(cell, bg_color)
            add_cell_text(cell, text, bold=True, size=9, color="FFFFFF",
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def set_table_borders(table, color="BDC3C7", sz="4"):
        """Set borders for entire table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    # ══════════════════════════════════════════════════════════════════════════
    # PAGE 1: MAIN CHARACTER SHEET
    # ══════════════════════════════════════════════════════════════════════════

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("ЛИСТ ПЕРСОНАЖА D&D 5e")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_HEADER_BG)

    # ── Basic Info Table ──
    add_section_header(doc, "ОСНОВНАЯ ИНФОРМАЦИЯ")

    t = doc.add_table(rows=3, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    # Row 0
    make_header_row(t, 0, ["Имя персонажа", "Класс и уровень", "Предыстория", "Имя игрока"])
    add_cell_text(t.cell(1, 0), character.name, size=10, bold=True)
    add_cell_text(t.cell(1, 1), f"{character.class_name} {character.level} ур.", size=10, bold=True)
    add_cell_text(t.cell(1, 2), character.background, size=10)
    add_cell_text(t.cell(1, 3), character.player_name, size=10)

    # Row 2
    make_header_row(t, 2, ["Раса", "Мировоззрение", "Очки опыта", "Вдохновение"])
    set_cell_shading(t.cell(2, 0), COLOR_LIGHT_BG)
    set_cell_shading(t.cell(2, 1), COLOR_LIGHT_BG)
    set_cell_shading(t.cell(2, 2), COLOR_LIGHT_BG)
    set_cell_shading(t.cell(2, 3), COLOR_LIGHT_BG)
    # Merge row 2 data into row 2 (reuse headers with data)
    # Actually, let's add data below headers
    # We need a 4-row table: headers + data + headers + data

    # Let me redo this with a proper 4-row table
    doc._body._body.remove(t._tbl)

    t = doc.add_table(rows=4, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["Имя персонажа", "Класс и уровень", "Предыстория", "Имя игрока"])
    add_cell_text(t.cell(1, 0), character.name, size=10, bold=True)
    add_cell_text(t.cell(1, 1), f"{character.class_name} {character.level} ур.", size=10, bold=True)
    add_cell_text(t.cell(1, 2), character.background, size=10)
    add_cell_text(t.cell(1, 3), character.player_name, size=10)

    make_header_row(t, 2, ["Раса", "Мировоззрение", "Очки опыта", "Вдохновение"],
                    bg_color=COLOR_SUBHEADER_BG)
    add_cell_text(t.cell(3, 0), character.race, size=10)
    add_cell_text(t.cell(3, 1), character.alignment, size=10)
    add_cell_text(t.cell(3, 2), str(character.experience_points), size=10)
    inspiration_text = "Да" if character.inspiration else "Нет"
    add_cell_text(t.cell(3, 3), inspiration_text, size=10)

    # Set column widths
    for row in t.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(4.5)

    # ── Ability Scores & Combat Stats ──
    add_section_header(doc, "ХАРАКТЕРИСТИКИ И БОЕВЫЕ ПАРАМЕТРЫ")

    # Proficiency bonus
    prof_bonus = character.get_proficiency_bonus()

    t = doc.add_table(rows=8, cols=7)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    # Header row
    make_header_row(t, 0, ["Характеристика", "Значение", "Бонус", "Спасбросок", "Владе-ние", "Бонус спасброска", ""])
    # Merge last column for combat stats label
    t.cell(0, 6).merge(t.cell(0, 6))
    set_cell_shading(t.cell(0, 6), COLOR_HEADER_BG)
    add_cell_text(t.cell(0, 6), "Боевые параметры", bold=True, size=9, color="FFFFFF",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Ability rows
    for i, abbr in enumerate(ABILITY_NAMES):
        row = i + 1
        total_score = character.get_total_ability_score(abbr)
        mod = character.get_modifier(abbr)
        saving = character.get_saving_throw(abbr)
        is_prof = character.saving_throw_proficiencies.get(abbr, False)

        add_cell_text(t.cell(row, 0), f"{ABILITY_FULL_NAMES[abbr]} ({abbr})", bold=True, size=9)
        add_cell_text(t.cell(row, 1), str(total_score), size=10, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(t.cell(row, 2), format_modifier(mod), size=10, bold=True,
                      color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(t.cell(row, 3), format_modifier(saving), size=10,
                      color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        prof_mark = "●" if is_prof else "○"
        add_cell_text(t.cell(row, 4), prof_mark, size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

        if is_prof:
            set_cell_shading(t.cell(row, 3), COLOR_LIGHT_BG)

        # Combat stats in right columns
        if row == 1:
            add_cell_text(t.cell(row, 5), "Бонус мастерства:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), format_modifier(prof_bonus), size=10, bold=True,
                          color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 2:
            ac_val = character.armor_class if character.armor_class is not None else 10 + character.get_modifier('ЛОВ')
            add_cell_text(t.cell(row, 5), "КД:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), str(ac_val), size=10, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 3:
            add_cell_text(t.cell(row, 5), "Инициатива:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), format_modifier(character.get_initiative()),
                          size=10, bold=True, color=COLOR_CALCULATED,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 4:
            add_cell_text(t.cell(row, 5), "Скорость:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), f"{character.speed} фт.", size=10,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 5:
            add_cell_text(t.cell(row, 5), "Макс. хитов:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), str(character.get_hp_max()), size=10, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 6:
            add_cell_text(t.cell(row, 5), "Текущие хиты:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), str(character.hp_current), size=10,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)
        elif row == 7:
            add_cell_text(t.cell(row, 5), "Временные хиты:", bold=True, size=8)
            add_cell_text(t.cell(row, 6), str(character.hp_temp), size=10,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)

    # Column widths for abilities table
    for row in t.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(1.5)
        row.cells[3].width = Cm(2.0)
        row.cells[4].width = Cm(1.5)
        row.cells[5].width = Cm(3.0)
        row.cells[6].width = Cm(3.0)

    # ── Hit Dice and Death Saves ──
    add_section_header(doc, "КОСТЬ ХИТОВ И СПАСБРОСКИ ОТ СМЕРТИ")

    t = doc.add_table(rows=2, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["Кость хитов", "Успехи", "", "Провалы", "", ""])
    add_cell_text(t.cell(1, 0), character.hit_dice, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    success_marks = "● " * character.death_save_successes + "○ " * (3 - character.death_save_successes)
    failure_marks = "● " * character.death_save_failures + "○ " * (3 - character.death_save_failures)

    t.cell(1, 1).merge(t.cell(1, 2))
    add_cell_text(t.cell(1, 1), success_marks.strip(), size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    t.cell(1, 3).merge(t.cell(1, 5))
    add_cell_text(t.cell(1, 3), failure_marks.strip(), size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Skills ──
    add_section_header(doc, "НАВЫКИ")

    # Skills in 2 columns
    skills_left = ALL_SKILLS[:9]
    skills_right = ALL_SKILLS[9:]

    t = doc.add_table(rows=10, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    # Header
    make_header_row(t, 0, ["Влад.", "Навык", "Бонус", "Влад.", "Навык", "Бонус"])

    for i in range(9):
        row = i + 1
        # Left column
        skill_l = skills_left[i]
        is_prof_l = character.skill_proficiencies.get(skill_l, False)
        is_expert_l = character.skill_expertise.get(skill_l, False)
        bonus_l = character.get_skill_bonus(skill_l)
        ability_l = SKILL_MAP[skill_l]

        prof_mark_l = "●●" if is_expert_l else ("●" if is_prof_l else "○")
        add_cell_text(t.cell(row, 0), prof_mark_l, size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(t.cell(row, 1), f"{skill_l} ({ability_l})", size=8)
        add_cell_text(t.cell(row, 2), format_modifier(bonus_l), size=9, bold=True,
                      color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        if is_prof_l:
            set_cell_shading(t.cell(row, 2), COLOR_LIGHT_BG)

        # Right column
        skill_r = skills_right[i]
        is_prof_r = character.skill_proficiencies.get(skill_r, False)
        is_expert_r = character.skill_expertise.get(skill_r, False)
        bonus_r = character.get_skill_bonus(skill_r)
        ability_r = SKILL_MAP[skill_r]

        prof_mark_r = "●●" if is_expert_r else ("●" if is_prof_r else "○")
        add_cell_text(t.cell(row, 3), prof_mark_r, size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(t.cell(row, 4), f"{skill_r} ({ability_r})", size=8)
        add_cell_text(t.cell(row, 5), format_modifier(bonus_r), size=9, bold=True,
                      color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        if is_prof_r:
            set_cell_shading(t.cell(row, 5), COLOR_LIGHT_BG)

    # Passive perception note
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pp.paragraph_format.space_before = Pt(4)
    run = pp.add_run(f"Пассивная Мудрость (Внимательность): ")
    run.font.size = Pt(9)
    run.font.bold = True
    run = pp.add_run(str(character.get_passive_perception()))
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_CALCULATED)

    # ── Attacks & Spellcasting ──
    add_section_header(doc, "АТАКИ И ЗАКЛИНАНИЯ")

    num_attacks = max(3, len(character.attacks))
    t = doc.add_table(rows=num_attacks + 1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["Название", "Бонус атаки", "Урон / Вид"])

    for i, atk in enumerate(character.attacks):
        add_cell_text(t.cell(i + 1, 0), atk.name, size=9)
        add_cell_text(t.cell(i + 1, 1), atk.attack_bonus, size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(t.cell(i + 1, 2), atk.damage_and_type, size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Fill empty rows
    for i in range(len(character.attacks), num_attacks):
        for j in range(3):
            add_cell_text(t.cell(i + 1, j), "", size=9)
            set_cell_shading(t.cell(i + 1, j), COLOR_VERY_LIGHT_BG)

    # ── Currency ──
    add_section_header(doc, "ВАЛЮТА")

    t = doc.add_table(rows=2, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["ММ (медь)", "СМ (серебро)", "ЭМ (электрум)", "ЗМ (золото)", "ПМ (платина)"])
    add_cell_text(t.cell(1, 0), str(character.cp), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 1), str(character.sp), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 2), str(character.ep), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 3), str(character.gp), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 4), str(character.pp), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Personality Section ──
    add_section_header(doc, "ЛИЧНОСТЬ")

    personality_items = [
        ("Черты характера", character.personality_traits),
        ("Идеалы", character.ideals),
        ("Привязанности", character.bonds),
        ("Слабости", character.flaws),
    ]

    t = doc.add_table(rows=len(personality_items), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    for i, (label, value) in enumerate(personality_items):
        set_cell_shading(t.cell(i, 0), COLOR_LIGHT_BG)
        add_cell_text(t.cell(i, 0), label, bold=True, size=9)
        add_cell_text(t.cell(i, 1), value if value else "", size=9)
        t.cell(i, 0).width = Cm(3.5)
        t.cell(i, 1).width = Cm(13.0)

    # ── Other Proficiencies & Languages ──
    add_section_header(doc, "ПРОЧИЕ ВЛАДЕНИЯ И ЯЗЫКИ")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.other_proficiencies_languages, size=9)

    # ── Equipment ──
    add_section_header(doc, "СНАРЯЖЕНИЕ")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.equipment, size=9)

    # ── Features & Traits ──
    add_section_header(doc, "УМЕНИЯ И ОСОБЕННОСТИ")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.features_traits, size=9)

    # ══════════════════════════════════════════════════════════════════════════
    # PAGE 2: CHARACTER DETAILS
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ДЕТАЛИ ПЕРСОНАЖА")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_HEADER_BG)

    # ── Physical Description ──
    add_section_header(doc, "ФИЗИЧЕСКОЕ ОПИСАНИЕ")

    t = doc.add_table(rows=2, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["Возраст", "Рост", "Вес", "Глаза", "Кожа", "Волосы"])
    add_cell_text(t.cell(1, 0), character.age, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 1), character.height, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 2), character.weight, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 3), character.eyes, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 4), character.skin, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 5), character.hair, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Character Appearance ──
    add_section_header(doc, "ВНЕШНОСТЬ ПЕРСОНАЖА")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.appearance, size=9)

    # ── Allies & Organizations ──
    add_section_header(doc, "СОЮЗНИКИ И ОРГАНИЗАЦИИ")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.allies_organizations, size=9)

    # ── Additional Features & Traits ──
    add_section_header(doc, "ДОПОЛНИТЕЛЬНЫЕ УМЕНИЯ И ОСОБЕННОСТИ")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.additional_features_traits, size=9)

    # ── Character Backstory ──
    add_section_header(doc, "ПРЕДЫСТОРИЯ ПЕРСОНАЖА")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.backstory, size=9)

    # ── Treasure ──
    add_section_header(doc, "СОКРОВИЩА")

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    add_cell_text(t.cell(0, 0), character.treasure, size=9)

    # ══════════════════════════════════════════════════════════════════════════
    # PAGE 3: SPELLCASTING
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ЗАКЛИНАНИЯ")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_HEADER_BG)

    # ── Spellcasting Stats ──
    add_section_header(doc, "ПАРАМЕТРЫ ЗАКЛИНАТЕЛЯ")

    spell_ability_full = ABILITY_FULL_NAMES.get(character.spellcasting_ability, "")
    spell_ability_mod = character.get_spellcasting_ability_modifier()

    t = doc.add_table(rows=2, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    make_header_row(t, 0, ["Класс заклинателя", "Характеристика", "Сложность спасения", "Бонус атаки закл.", "Модификатор хар-ки"])
    add_cell_text(t.cell(1, 0), character.spellcasting_class, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 1), f"{spell_ability_full} ({character.spellcasting_ability})" if character.spellcasting_ability else "—",
                  size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 2), str(character.get_spell_save_dc()),
                  size=10, bold=True, color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 3), format_modifier(character.get_spell_attack_bonus()),
                  size=10, bold=True, color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(t.cell(1, 4), format_modifier(spell_ability_mod),
                  size=10, bold=True, color=COLOR_CALCULATED, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Spell Slots ──
    add_section_header(doc, "ЯЧЕЙКИ ЗАКЛИНАНИЙ")

    # Count how many levels have slots
    slot_levels = []
    for lvl in range(1, 10):
        slot_info = character.spell_slots.get(lvl)
        if slot_info and slot_info.total_slots > 0:
            slot_levels.append(lvl)

    if slot_levels:
        t = doc.add_table(rows=3, cols=len(slot_levels))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)

        # Headers
        for i, lvl in enumerate(slot_levels):
            add_cell_text(t.cell(0, i), f"{lvl} ур.", bold=True, size=9,
                          color="FFFFFF", alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(t.cell(0, i), COLOR_HEADER_BG)

        # Total slots
        for i, lvl in enumerate(slot_levels):
            slot_info = character.spell_slots.get(lvl, SpellSlotInfo(lvl))
            add_cell_text(t.cell(1, i), str(slot_info.total_slots), size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(t.cell(1, i), COLOR_LIGHT_BG)

        # Expended
        for i, lvl in enumerate(slot_levels):
            slot_info = character.spell_slots.get(lvl, SpellSlotInfo(lvl))
            add_cell_text(t.cell(2, i), str(slot_info.expended_slots), size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        t = doc.add_table(rows=2, cols=9)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        for i in range(9):
            add_cell_text(t.cell(0, i), f"{i+1} ур.", bold=True, size=8,
                          color="FFFFFF", alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(t.cell(0, i), COLOR_HEADER_BG)
            add_cell_text(t.cell(1, i), "0", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Cantrips ──
    add_section_header(doc, "ЗАГОВОРЫ (0 уровень)")

    if character.cantrips:
        t = doc.add_table(rows=len(character.cantrips), cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        make_header_row(t, 0, ["#", "Название"])
        for i, spell_name in enumerate(character.cantrips):
            add_cell_text(t.cell(i, 0), str(i + 1), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell_text(t.cell(i, 1), spell_name, size=9)
    else:
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        add_cell_text(t.cell(0, 0), "(нет заговоров)", size=9, italic=True,
                      color=COLOR_BORDER)

    # ── Spells by Level ──
    for lvl in range(1, 10):
        spells = character.spells_by_level.get(lvl, [])
        if not spells:
            continue

        add_section_header(doc, f"ЗАКЛИНАНИЯ {lvl} УРОВНЯ")

        t = doc.add_table(rows=len(spells), cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        make_header_row(t, 0, ["Подг.", "Название", "Прим."])
        for i, spell in enumerate(spells):
            prep = "●" if spell.prepared else "○"
            add_cell_text(t.cell(i, 0), prep, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell_text(t.cell(i, 1), spell.name, size=9)
            add_cell_text(t.cell(i, 2), "", size=9)

    # ── Legend ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Обозначения: ")
    run.font.size = Pt(8)
    run.font.bold = True
    run = p.add_run("● = владение / подготовлено  ○ = нет владения / не подготовлено  ●● = экспертиза  ")
    run.font.size = Pt(8)
    run = p.add_run("Фиолетовый цвет = авторасчёт")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLOR_CALCULATED)

    # ── Save ──
    doc.save(output_path)
    return output_path


# ─── Interactive Character Builder ────────────────────────────────────────────

def interactive_build() -> Character:
    """Build a character interactively via CLI."""
    print("=" * 60)
    print("  ГЕНЕРАТОР ЛИСТА ПЕРСОНАЖА D&D 5e")
    print("  Интерактивный ввод данных персонажа")
    print("=" * 60)
    print()

    char = Character()
    char.init_skill_defaults()

    def ask(prompt: str, default: str = "") -> str:
        suffix = f" [{default}]" if default else ""
        val = input(f"{prompt}{suffix}: ").strip()
        return val if val else default

    def ask_int(prompt: str, default: int = 0) -> int:
        val = ask(prompt, str(default))
        try:
            return int(val)
        except ValueError:
            return default

    def ask_bool(prompt: str, default: bool = False) -> bool:
        val = ask(f"{prompt} (д/н)", "д" if default else "н").lower()
        return val in ('д', 'y', 'yes', 'да')

    # Basic info
    print("\n── Основная информация ──")
    char.name = ask("Имя персонажа")
    char.class_name = ask("Класс")
    char.level = ask_int("Уровень", 1)
    char.background = ask("Предыстория")
    char.player_name = ask("Имя игрока")
    char.race = ask("Раса")
    char.alignment = ask("Мировоззрение")
    char.experience_points = ask_int("Очки опыта", 0)
    char.inspiration = ask_bool("Вдохновение?")

    # Ability scores
    print("\n── Характеристики (базовые значения без расовых бонусов) ──")
    for abbr in ABILITY_NAMES:
        char.ability_scores[abbr] = ask_int(f"  {ABILITY_FULL_NAMES[abbr]} ({abbr})", 10)

    print("\n── Расовые/прочие бонусы к характеристикам ──")
    for abbr in ABILITY_NAMES:
        char.ability_bonuses[abbr] = ask_int(f"  Бонус к {ABILITY_FULL_NAMES[abbr]} ({abbr})", 0)

    # Saving throw proficiencies
    print("\n── Владение спасбросками ──")
    for abbr in ABILITY_NAMES:
        char.saving_throw_proficiencies[abbr] = ask_bool(f"  {ABILITY_FULL_NAMES[abbr]} ({abbr})?")

    # Skill proficiencies
    print("\n── Владение навыками ──")
    for skill in ALL_SKILLS:
        char.skill_proficiencies[skill] = ask_bool(f"  {skill} ({SKILL_MAP[skill]})?")

    # Skill expertise
    print("\n── Экспертиза навыков (двойной бонус мастерства) ──")
    for skill in ALL_SKILLS:
        if char.skill_proficiencies[skill]:
            char.skill_expertise[skill] = ask_bool(f"  Экспертиза: {skill}?")

    # Combat
    print("\n── Боевые параметры ──")
    ac_input = ask("КД (пустое = 10 + мод. ЛОВ)")
    char.armor_class = int(ac_input) if ac_input else None

    init_input = ask("Инициатива (пустое = мод. ЛОВ)")
    char.initiative_override = int(init_input) if init_input else None

    char.speed = ask_int("Скорость (футы)", 30)
    char.hp_max = ask_int("Максимум хитов (0 = ввод вручную)", 0) or None
    char.hp_current = ask_int("Текущие хиты", 0)
    char.hp_temp = ask_int("Временные хиты", 0)
    char.hit_dice = ask("Кость хитов (напр. 1d10)", "")

    # Death saves
    print("\n── Спасброски от смерти ──")
    char.death_save_successes = ask_int("Успехи (0-3)", 0)
    char.death_save_failures = ask_int("Провалы (0-3)", 0)

    # Attacks
    print("\n── Атаки ──")
    num_attacks = ask_int("Количество атак", 0)
    for i in range(num_attacks):
        print(f"  Атака {i + 1}:")
        atk = Attack()
        atk.name = ask("    Название")
        atk.attack_bonus = ask("    Бонус атаки")
        atk.damage_and_type = ask("    Урон / Вид")
        char.attacks.append(atk)

    # Currency
    print("\n── Валюта ──")
    char.cp = ask_int("ММ (медь)", 0)
    char.sp = ask_int("СМ (серебро)", 0)
    char.ep = ask_int("ЭМ (электрум)", 0)
    char.gp = ask_int("ЗМ (золото)", 0)
    char.pp = ask_int("ПМ (платина)", 0)

    # Personality
    print("\n── Личность ──")
    char.personality_traits = ask("Черты характера")
    char.ideals = ask("Идеалы")
    char.bonds = ask("Привязанности")
    char.flaws = ask("Слабости")

    # Other proficiencies & languages
    print("\n── Прочие владения и языки ──")
    char.other_proficiencies_languages = ask("Владения и языки")

    # Features & traits
    print("\n── Умения и особенности ──")
    char.features_traits = ask("Умения и особенности")

    # Equipment
    print("\n── Снаряжение ──")
    char.equipment = ask("Снаряжение")

    # Character details (page 2)
    print("\n── Детали персонажа (страница 2) ──")
    char.age = ask("Возраст")
    char.height = ask("Рост")
    char.weight = ask("Вес")
    char.eyes = ask("Глаза")
    char.skin = ask("Кожа")
    char.hair = ask("Волосы")
    char.appearance = ask("Внешность персонажа")
    char.allies_organizations = ask("Союзники и организации")
    char.additional_features_traits = ask("Дополнительные умения и особенности")
    char.backstory = ask("Предыстория персонажа")
    char.treasure = ask("Сокровища")

    # Spellcasting (page 3)
    print("\n── Заклинания (страница 3) ──")
    char.spellcasting_class = ask("Класс заклинателя")
    if char.spellcasting_class:
        print("  Характеристика заклинаний:")
        for abbr in ABILITY_NAMES:
            print(f"    {abbr} - {ABILITY_FULL_NAMES[abbr]}")
        char.spellcasting_ability = ask("  Характеристика (СИЛ/ЛОВ/ТЕЛ/ИНТ/МДР/ХАР)")

        # Spell slots
        print("\n  Ячейки заклинаний:")
        for lvl in range(1, 10):
            total = ask_int(f"    Уровень {lvl} - Всего ячеек", 0)
            if total > 0:
                expended = ask_int(f"    Уровень {lvl} - Потрачено", 0)
                char.spell_slots[lvl] = SpellSlotInfo(lvl, total, expended)

        # Cantrips
        print("\n  Заговоры:")
        num_cantrips = ask_int("  Количество заговоров", 0)
        for i in range(num_cantrips):
            char.cantrips.append(ask(f"    Заговор {i + 1}"))

        # Spells by level
        print("\n  Заклинания по уровням:")
        for lvl in range(1, 10):
            num_spells = ask_int(f"  Количество заклинаний {lvl} уровня", 0)
            if num_spells > 0:
                char.spells_by_level[lvl] = []
                for i in range(num_spells):
                    name = ask(f"    Заклинание {i + 1}")
                    prepared = ask_bool(f"    Подготовлено: {name}?")
                    char.spells_by_level[lvl].append(SpellEntry(name, prepared))

    return char


# ─── JSON-based Character Builder ────────────────────────────────────────────

def create_example_character() -> Character:
    """Create an example character for testing."""
    char = Character()
    char.init_skill_defaults()

    char.name = "Торин Каменный Щит"
    char.class_name = "Воин"
    char.level = 5
    char.background = "Солдат"
    char.player_name = "Алексей"
    char.race = "Дворф (Горный)"
    char.alignment = "Законно-добрый"
    char.experience_points = 6500
    char.inspiration = True

    # Ability scores (base)
    char.ability_scores = {
        'СИЛ': 16, 'ЛОВ': 10, 'ТЕЛ': 14,
        'ИНТ': 8, 'МДР': 12, 'ХАР': 10
    }

    # Racial bonuses (Mountain Dwarf: +2 STR, +2 CON)
    char.ability_bonuses = {
        'СИЛ': 2, 'ЛОВ': 0, 'ТЕЛ': 2,
        'ИНТ': 0, 'МДР': 0, 'ХАР': 0
    }

    # Saving throw proficiencies (Fighter: STR, CON)
    char.saving_throw_proficiencies = {
        'СИЛ': True, 'ЛОВ': False, 'ТЕЛ': True,
        'ИНТ': False, 'МДР': False, 'ХАР': False
    }

    # Skill proficiencies (Fighter + Soldier background)
    char.skill_proficiencies = {
        'Акробатика': False, 'Анализ': False, 'Атлетика': True,
        'Внимательность': True, 'Выживание': False, 'Выступление': False,
        'Запугивание': True, 'История': False, 'Ловкость рук': False,
        'Магия': False, 'Медицина': False, 'Обман': False,
        'Природа': False, 'Проницательность': False, 'Религия': False,
        'Скрытность': False, 'Убеждение': False, 'Уход за животными': False,
    }
    char.skill_expertise = {k: False for k in ALL_SKILLS}

    # Combat
    char.armor_class = 18  # Chain mail
    char.speed = 25  # Dwarf
    char.hp_max = 49
    char.hp_current = 49
    char.hp_temp = 0
    char.hit_dice = "5d10"

    # Death saves
    char.death_save_successes = 0
    char.death_save_failures = 0

    # Attacks
    char.attacks = [
        Attack("Боевой топор +1", "+7", "1d8+4 рубящий"),
        Attack("Ручной арбалет", "+5", "1d6 колющий"),
    ]

    # Currency
    char.cp = 47
    char.sp = 12
    char.ep = 0
    char.gp = 85
    char.pp = 2

    # Personality
    char.personality_traits = "Я всегда планирую заранее, даже в самых простых ситуациях. Я могу решать проблемы решительно и быстро."
    char.ideals = "Честь — превыше всего. Моё слово — закон, и я не нарушу данное обещание ни при каких обстоятельствах."
    char.bonds = "Я сражаюсь за тех, кто не может постоять за себя. Мой клан — моя семья, и я защищу каждого его члена до последней капли крови."
    char.flaws = "Я слишком упрям и редко меняю своё решение, даже когда очевидно, что ошибаюсь. Это может доводить моих спутников до ярости."

    # Other proficiencies & languages
    char.other_proficiencies_languages = (
        "Владение: Все доспехи, щиты, простое и воинское оружие\n"
        "Инструменты: Набор кузнеца\n"
        "Языки: Общий, Дворфийский"
    )

    # Features & traits
    char.features_traits = (
        "Второе дыхание: Восстановление 1d10+5 хитов (1/короткий отдых)\n"
        "Действенный удар: Доп. атака бонусным действием\n"
        "Улучшение характеристики: СИЛ +2 (уровень 4)\n"
        "Дворфья выносливость: Устойчивость к яду\n"
        "Дворфья боевая тренировка: Владение боевым топором, ручным арбалетом, молотом, молотом войны\n"
        "Знание камня: Преимущество на проверки Истории по каменной кладке"
    )

    # Equipment
    char.equipment = (
        "Кольчуга, Щит, Боевой топор +1, Ручной арбалет (20 болтов), "
        "Набор путешественника, Символ клана, Игральные кости, "
        "Набор шахтёра, Рюкзак, Вёрёвка (50 фт.), Факелы (10), "
        "Сухой паёк (10 дней), Фляга воды"
    )

    # Character details
    char.age = "62"
    char.height = "4'5\" (135 см)"
    char.weight = "170 фунтов (77 кг)"
    char.eyes = "Карие"
    char.skin = "Загорелая, обветренная"
    char.hair = "Рыжая, заплетённая в косы"
    char.appearance = (
        "Торин — коренастый дворф с широкой грудью и мощными руками кузнеца. "
        "Его рыжая борода заплетена в две толстые косы, украшенные медными кольцами — "
        "по одному за каждую значимую победу. На левом глазу — старый шрам от битвы с орками. "
        "На правом плече — вытатуированный символ клана Каменного Щита: молот, "
        "разбивающий гору. Носит потёртую кольчугу, за которой тщательно ухаживает, "
        "и всегда начищает свой боевой топор до блеска."
    )
    char.allies_organizations = (
        "Клан Каменного Щита — древний дворфийский клан из Цитадели Адбар. "
        "Капитан Гаррик — бывший командир в армии Лорда, который обучил Торина тактике. "
        "Гильдия кузнецов Миробара — обеспечивает доступ к редким металлам."
    )
    char.additional_features_traits = (
        "Новая черта на 5 уровне: Дополнительная атака — при действии Атаки "
        "можно атаковать дважды вместо одного раза."
    )
    char.backstory = (
        "Торин родился в глубинах Цитадели Адбар, в знаменитом клане кузнецов Каменного Щита. "
        "С детства он помогал отцу в кузнице, но вместо того, чтобы стать мастером ковки, "
        "почувствовал зов битвы. Когда орды орков напали на цитадель, восемнадцатилетний Торин "
        "взял в руки боевой топор вместо молота и встал на защиту родного города.\n\n"
        "После разгрома орков Торин поступил на службу к местному лорду и провёл двадцать лет "
        "в армии, дослужившись до сержанта. Он участвовал в десятках сражений и заработал "
        "почтительное прозвище «Каменный Щит» — за то, что в бою стоял нерушимо, как скала.\n\n"
        "Однако мирная жизнь не для дворфа-воина. Когда он узнал о древнем артефакте, "
        "способном защитить его клан от надвигающейся угрозы из Подземья, Торин покинул службу "
        "и отправился в странствие. Теперь он ищет этот артефакт, попутно помогая тем, "
        "кто не может постоять за себя — ведь именно этому его учил капитан Гаррик."
    )
    char.treasure = (
        "Кольцо защиты +1, Свинцовый слиток с рунами клана, "
        "Старинная карта Подземья, Неогранённый рубин (стоимость 250 зм)"
    )

    # Spellcasting (not a spellcaster in this example)
    char.spellcasting_class = ""
    char.spellcasting_ability = ""

    return char


def create_example_spellcaster() -> Character:
    """Create an example spellcaster character for testing."""
    char = Character()
    char.init_skill_defaults()

    char.name = "Элара Звёздный Ветер"
    char.class_name = "Волшебник"
    char.level = 5
    char.background = "Мудрец"
    char.player_name = "Мария"
    char.race = "Высший эльф"
    char.alignment = "Хаотично-добрый"
    char.experience_points = 6500
    char.inspiration = False

    char.ability_scores = {
        'СИЛ': 8, 'ЛОВ': 14, 'ТЕЛ': 12,
        'ИНТ': 17, 'МДР': 12, 'ХАР': 10
    }
    char.ability_bonuses = {
        'СИЛ': 0, 'ЛОВ': 2, 'ТЕЛ': 0,
        'ИНТ': 1, 'МДР': 0, 'ХАР': 0
    }

    char.saving_throw_proficiencies = {
        'СИЛ': False, 'ЛОВ': False, 'ТЕЛ': False,
        'ИНТ': True, 'МДР': True, 'ХАР': False
    }

    char.skill_proficiencies = {
        'Акробатика': False, 'Анализ': True, 'Атлетика': False,
        'Внимательность': False, 'Выживание': False, 'Выступление': False,
        'Запугивание': False, 'История': True, 'Ловкость рук': False,
        'Магия': True, 'Медицина': False, 'Обман': False,
        'Природа': False, 'Проницательность': True, 'Религия': True,
        'Скрытность': False, 'Убеждение': False, 'Уход за животными': False,
    }
    char.skill_expertise = {k: False for k in ALL_SKILLS}

    char.armor_class = 12  # Mage armor
    char.speed = 30
    char.hp_max = 28
    char.hp_current = 28
    char.hp_temp = 0
    char.hit_dice = "5d6"

    char.attacks = [
        Attack("Огненный снаряд", "+7", "1d10+4 огонь"),
    ]

    char.cp = 15
    char.sp = 8
    char.gp = 42
    char.pp = 0

    char.personality_traits = "Я одержима знаниями и всегда ищу новые заклинания и тайны."
    char.ideals = "Знание должно быть свободным. Я делюсь открытиями со всеми, кто ищет истину."
    char.bonds = "Древний гримуар моей наставницы — моя самая ценная вещь. Я храню его как зеницу ока."
    char.flaws = "Я легко отвлекаюсь на интересные магические феномены, даже в опасных ситуациях."

    char.other_proficiencies_languages = (
        "Владение: Кинжалы, дротики, пращи, посохи\n"
        "Инструменты: Нет\n"
        "Языки: Общий, Эльфийский, Драконий, Инфернальный"
    )

    char.features_traits = (
        "Восстановление дуги: Восстановление ячеек заклинаний (1/приключение)\n"
        "Зачарование: Улучшение заклинаний за 2 дополнительные ячейки\n"
        "Улучшение характеристики: ИНТ +2 (уровень 4)\n"
        "Эльфийское наследие: Тёмное зрение 60 фт., Чувство фей, Транса"
    )

    char.equipment = (
        "Посох, Компонентная сумка, Книга заклинаний, "
        "Набор учёного, Фонарь, Чернила и перо, "
        "Сухой паёк (5 дней), Фляга воды"
    )

    char.age = "125"
    char.height = "5'7\" (170 см)"
    char.weight = "130 фунтов (59 кг)"
    char.eyes = "Серебристые"
    char.skin = "Бледная, с лёгким серебристым отливом"
    char.hair = "Чёрная, с серебряными прядями"
    char.appearance = (
        "Элара — изящная эльфийка с серебристыми глазами, которые, кажется, "
        "светятся в темноте. Её длинные чёрные волосы украшены серебряными прядями — "
        "не от старости, а от магического воздействия древнего гримуара. "
        "Она носит тёмно-синюю мантию, расшитую серебряными рунами, "
        "и всегда держит посох из белого ясеня, увенчанный лунным камнем."
    )
    char.allies_organizations = (
        "Академия Высокой Магии Сильвермун — Альма-матер.\n"
        "Мастер Эллион — наставница, исчезнувшая при загадочных обстоятельствах."
    )
    char.backstory = (
        "Элара родилась в эльфийском городе Сильвермун, в семье библиотекарей. "
        "С раннего детства она проводила дни среди древних фолиантов, впитывая знания "
        "о магических традициях забытых цивилизаций. В сто лет она поступила "
        "в Академию Высокой Магии, где проявила выдающийся талант к школе Воплощения.\n\n"
        "Её наставница, мастер Эллион, обучала её не только заклинаниям, но и "
        "исследованию магических аномалий. Когда Эллион загадочно исчезла, оставив "
        "лишь свой гримуар и записку с координатами древнего руинного комплекса, "
        "Элара отправилась на поиски, движимая верностью и жаждой знаний.\n\n"
        "С тех пор она странствует по миру, раскрывая тайны магии и разыскивая "
        "следы своей наставницы."
    )
    char.treasure = "Гримуар Эллион, Кристалл скрайинга (расходуемый), Свиток невидимости"

    # Spellcasting
    char.spellcasting_class = "Волшебник"
    char.spellcasting_ability = "ИНТ"

    # Spell slots for level 5 wizard
    char.spell_slots = {
        1: SpellSlotInfo(1, 4, 0),
        2: SpellSlotInfo(2, 3, 0),
        3: SpellSlotInfo(3, 2, 0),
    }

    # Cantrips
    char.cantrips = [
        "Огненный снаряд",
        "Маленькая иллюзия",
        "Престидижитация",
        "Луч морозa",
    ]

    # Spells
    char.spells_by_level = {
        1: [
            SpellEntry("Опознание", True),
            SpellEntry("Магический снаряд", True),
            SpellEntry("Щит", True),
            SpellEntry("Доспех мага", True),
            SpellEntry("Обнаружение магии", True),
            SpellEntry("Волна грома", False),
        ],
        2: [
            SpellEntry("Невидимость", True),
            SpellEntry("Паутина", True),
            SpellEntry("Обнаружение мыслей", False),
            SpellEntry("Туманный шаг", True),
        ],
        3: [
            SpellEntry("Огненный шар", True),
            SpellEntry("Молния", True),
            SpellEntry("Полёт", False),
        ],
    }

    return char


# ─── JSON Import/Export ───────────────────────────────────────────────────────

def save_character_json(character: Character, path: str):
    """Save character to JSON file."""
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(character.to_dict(), f, ensure_ascii=False, indent=2)
    print(f"Персонаж сохранён в JSON: {path}")


def load_character_json(path: str) -> Character:
    """Load character from JSON file."""
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return Character.from_dict(data)


# ─── Main Entry Point ─────────────────────────────────────────────────────────

def main():
    """Main entry point."""
    import sys

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║   D&D 5e ГЕНЕРАТОР ЛИСТА ПЕРСОНАЖА                ║")
    print("║   Авторасчёт → Экспорт в DOCX                     ║")
    print("╚══════════════════════════════════════════════════════╝")
    print()
    print("Режимы работы:")
    print("  1 — Интерактивный ввод (пошаговый)")
    print("  2 — Пример: Воин 5 уровня (Торин)")
    print("  3 — Пример: Волшебник 5 уровня (Элара)")
    print("  4 — Загрузить из JSON файла")
    print()

    choice = input("Выберите режим [1-4]: ").strip()

    character = None

    if choice == "1":
        character = interactive_build()
    elif choice == "2":
        character = create_example_character()
        print(f"\nЗагружен пример: {character.name}")
    elif choice == "3":
        character = create_example_spellcaster()
        print(f"\nЗагружен пример: {character.name}")
    elif choice == "4":
        json_path = input("Путь к JSON файлу: ").strip()
        try:
            character = load_character_json(json_path)
            print(f"\nЗагружен персонаж: {character.name}")
        except Exception as e:
            print(f"Ошибка загрузки: {e}")
            sys.exit(1)
    else:
        print("Неверный выбор.")
        sys.exit(1)

    # Print calculated stats summary
    print("\n" + "=" * 60)
    print(f"  ПЕРСОНАЖ: {character.name}")
    print(f"  {character.class_name} {character.level} уровня | {character.race} | {character.alignment}")
    print("=" * 60)

    prof = character.get_proficiency_bonus()
    print(f"\n  Бонус мастерства: +{prof}")
    print(f"\n  Характеристики:")
    for abbr in ABILITY_NAMES:
        total = character.get_total_ability_score(abbr)
        mod = character.get_modifier(abbr)
        save = character.get_saving_throw(abbr)
        save_prof = "●" if character.saving_throw_proficiencies.get(abbr, False) else "○"
        print(f"    {ABILITY_FULL_NAMES[abbr]:14s} ({abbr}): {total:2d}  мод: {format_modifier(mod):3s}  спас: {format_modifier(save):3s} {save_prof}")

    print(f"\n  Навыки:")
    for skill in ALL_SKILLS:
        is_prof = character.skill_proficiencies.get(skill, False)
        is_expert = character.skill_expertise.get(skill, False)
        bonus = character.get_skill_bonus(skill)
        ability = SKILL_MAP[skill]
        mark = "●●" if is_expert else ("●" if is_prof else "○")
        print(f"    {mark} {skill:24s} ({ability}): {format_modifier(bonus)}")

    print(f"\n  Боевые:")
    ac = character.armor_class if character.armor_class is not None else 10 + character.get_modifier('ЛОВ')
    print(f"    КД: {ac}  Инициатива: {format_modifier(character.get_initiative())}  Скорость: {character.speed} фт.")
    print(f"    Хиты: {character.hp_current}/{character.get_hp_max()}")

    if character.spellcasting_class:
        print(f"\n  Заклинания:")
        print(f"    Класс: {character.spellcasting_class}")
        print(f"    Характеристика: {character.spellcasting_ability}")
        print(f"    Сложность спасения: {character.get_spell_save_dc()}")
        print(f"    Бонус атаки: {format_modifier(character.get_spell_attack_bonus())}")

    print(f"\n  Пассивная Внимательность: {character.get_passive_perception()}")

    # Generate DOCX
    output_dir = "/home/z/my-project/download"
    os.makedirs(output_dir, exist_ok=True)

    safe_name = character.name.replace(" ", "_").replace("/", "_") if character.name else "character"
    docx_path = os.path.join(output_dir, f"DnD5e_{safe_name}.docx")

    try:
        generate_docx(character, docx_path)
        print(f"\n✓ Лист персонажа сохранён: {docx_path}")
    except Exception as e:
        print(f"\n✗ Ошибка генерации DOCX: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Also save JSON
    json_path = os.path.join(output_dir, f"DnD5e_{safe_name}.json")
    save_character_json(character, json_path)

    print(f"\nГотово! Файлы:")
    print(f"  DOCX: {docx_path}")
    print(f"  JSON: {json_path}")


if __name__ == "__main__":
    main()
